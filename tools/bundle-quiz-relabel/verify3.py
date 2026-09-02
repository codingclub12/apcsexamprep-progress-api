import zipfile, re, sys
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import RGBColor
import plan
fails=[]
def chk(c,m):
    print(("  ok   " if c else "  FAIL ")+m)
    if not c: fails.append(m)

PAIRS=[("2.1","KEY"),("2.1","STUDENT"),("2.2","KEY"),("2.2","STUDENT")]
print("### A. every XML part is well-formed ###")
for l,v in PAIRS:
    z=zipfile.ZipFile("final/Lesson-%s-Quiz-%s.docx"%(l,v))
    bad=[]
    for n in z.namelist():
        if n.endswith("/"): continue
        try: ET.fromstring(z.read(n))
        except Exception as e: bad.append("%s: %s"%(n,e))
    chk(not bad,"%s-%s all %d XML parts parse %s"%(l,v,len([x for x in z.namelist() if not x.endswith('/')]),bad))

print("\n### B. nothing but word/document.xml was modified ###")
for l,v in PAIRS:
    a=zipfile.ZipFile("orig/%s-%s.docx"%(l,v)); b=zipfile.ZipFile("final/Lesson-%s-Quiz-%s.docx"%(l,v))
    diff=[n for n in a.namelist() if not n.endswith("/") and a.read(n)!=b.read(n)]
    chk(diff==["word/document.xml"],"%s-%s only document.xml differs (got %s)"%(l,v,diff))

print("\n### C. independent OPC parser (python-docx) opens each file ###")
for l,v in PAIRS:
    try:
        d=Document("final/Lesson-%s-Quiz-%s.docx"%(l,v))
        n=len(d.paragraphs)
        chk(n>0,"%s-%s opens cleanly via python-docx, %d paragraphs"%(l,v,n))
    except Exception as e:
        chk(False,"%s-%s FAILED to open: %s"%(l,v,e))

print("\n### D. purple + checkmark land on the right option (as the renderer sees it) ###")
for l in ("2.1","2.2"):
    d=Document("final/Lesson-%s-Quiz-KEY.docx"%l)
    marked={}
    for p in d.paragraphs:
        t=p.text.strip()
        m=re.match(r"^([A-D])\.\s",t)
        if not m: continue
        purple=any(r.font.color is not None and r.font.color.rgb==RGBColor(0x6B,0x21,0xA8)
                   for r in p.runs if r.text.strip())
        bold=any(r.bold for r in p.runs if r.text.strip() and r.text.strip() not in "ABCD.")
        if purple and "✓" in t:
            marked.setdefault(len(marked)+0,None)
    # walk questions in order
    seq=[]; cur=None
    for p in d.paragraphs:
        t=p.text.strip()
        if re.match(r"^\d+\.\s",t): cur=int(re.match(r"^(\d+)",t).group(1))
        m=re.match(r"^([A-D])\.\s",t)
        if m and cur:
            purple=any(r.font.color is not None and r.font.color.rgb==RGBColor(0x6B,0x21,0xA8)
                       for r in p.runs if r.text.strip())
            if purple and "✓" in t: seq.append((cur,m.group(1)))
    got="".join(x[1] for x in sorted(set(seq)))
    chk(got==plan.KEYS[l],"%s KEY: renderer-visible purple+checkmark sequence = %s"%(l,got))
    chk(len(set(q for q,_ in seq))==len(plan.KEYS[l]),"%s KEY: every question has exactly one marked option"%l)
    d2=Document("final/Lesson-%s-Quiz-STUDENT.docx"%l)
    anyp=any(r.font.color is not None and r.font.color.rgb==RGBColor(0x6B,0x21,0xA8)
             for p in d2.paragraphs for r in p.runs if re.match(r"^[A-D]\.\s",p.text.strip()))
    chk(not anyp,"%s STUDENT: no option is purple-marked"%l)
    chk("✓" not in "\n".join(p.text for p in d2.paragraphs),"%s STUDENT: no checkmarks leak"%l)

print("\n"+("ALL CHECKS PASS" if not fails else "%d FAILURES"%len(fails)))
sys.exit(1 if fails else 0)
