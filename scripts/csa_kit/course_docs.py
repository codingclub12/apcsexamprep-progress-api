"""
Course-level documents for the AP CSA teacher kit.

Start_Here.docx, How_To_Use_This_Course.docx, and Course_Resources/ with two
pacing guides. Modeled on what AP Cybersecurity already ships, because a
teacher who has used that course should find the same three documents in the
same places.

WHERE THE NUMBERS COME FROM, AND WHERE THEY DO NOT.

  Teaching days per unit are COUNTED from the authored kit content, one per
  Day<N> deck pair, so they cannot drift from what is in the folders.

  CED class-period bands are quoted from College Board's Course at a Glance,
  extracted verbatim to docs/csa-ced-course-at-a-glance.txt. Note the units
  differ: the CED counts 45-minute periods and this kit is built on 60-minute
  days, which is why 76 kit days is not short of ~99-105 CED periods. See
  docs/csa-deck-pacing-vs-ced.md for the arithmetic.

  Progress Check topic splits are the CED's own, all ten of them.

  ASSESSMENT DAYS ARE NOT INVENTED HERE. Tanner settled two of them on
  2026-08-24 and they are recorded: Unit 4 gets two test days, multiple choice
  then free response, and the mock exam runs across two days. Those are stated.
  A full 160-day calendar is NOT printed, because the Year Map that carries it
  is not in this repo and reproducing it from memory would be a guess wearing
  a table's clothes.
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x10, 0x24, 0x3A)
ACCENT = RGBColor(0x2C, 0x6B, 0xAF)
MUTED = RGBColor(0x5A, 0x6B, 0x7B)

TRADEMARK = ("AP is a trademark of the College Board, which was not involved in "
             "the production of, and does not endorse, this resource.")

# CED Course at a Glance. Class periods are 45 minutes in the CED's own words:
# "pacing is based on 45-minute class periods, meeting five days each week for
# a full academic year."
CED = {
    '1': ('Using Objects and Methods', '~32-34', '15-25%',
          [('Part 1', 'Topics 1.1-1.4'), ('Part 2', 'Topics 1.5-1.9'),
           ('Part 3', 'Topics 1.10-1.15')]),
    '2': ('Selection and Iteration', '~29-31', '25-35%',
          [('Part 1', 'Topics 2.1-2.6'), ('Part 2', 'Topics 2.7-2.12')]),
    '3': ('Class Creation', '~20-22', '10-18%',
          [('Part 1', 'Topics 3.1-3.4'), ('Part 2', 'Topics 3.5-3.9')]),
    '4': ('Data Collections', '~50-52', '30-40%',
          [('Part 1', 'Topics 4.1-4.5'), ('Part 2', 'Topics 4.6-4.10'),
           ('Part 3', 'Topics 4.11-4.17')]),
}


def _p(doc, text='', size=11, bold=False, color=None, space_after=6, font='Calibri'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = font
    if color is not None:
        r.font.color.rgb = color
    return p


def _h(doc, text, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.name = 'Cambria'
    r.font.color.rgb = NAVY
    return p


def _b(doc, text, size=10.5, indent=0.25):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    return p


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    r.font.name = 'Calibri'
    return t


def _foot(doc, label):
    _p(doc)
    _p(doc, TRADEMARK, size=7.5, color=MUTED, space_after=0)
    _p(doc, f'APCSExamPrep.com   {label}', size=7.5, color=MUTED, space_after=0)


def _title(doc, eyebrow, title, sub):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)
    _p(doc, eyebrow, size=10, bold=True, color=ACCENT, space_after=2)
    _p(doc, title, size=20, bold=True, color=NAVY, space_after=2)
    _p(doc, sub, size=9.5, color=MUTED, space_after=10)


def build_start_here(path, day_counts):
    doc = Document()
    _title(doc, 'AP COMPUTER SCIENCE A', 'Start Here',
           'A 60-second orientation: where everything lives, and the two '
           'documents to open first.')

    _h(doc, 'Open these first')
    _b(doc, 'How_To_Use_This_Course.docx. The teaching workflow: what every '
            'file is, how the three surfaces fit together, and the daily rhythm.')
    _b(doc, 'Course_Resources/. The full-year and block pacing guides.')

    _h(doc, 'How the course is organized')
    _p(doc, 'Four units, 53 topics, built against the 2025 Course and Exam '
            'Description. Every lesson folder contains:', size=10.5, space_after=4)
    _b(doc, 'Slide_Decks/. Day<N>_Deck_TEACHER.pptx with speaker notes, and '
            'Day<N>_Deck_STUDENT.pptx without them.')
    _b(doc, 'Guided_Notes/. Day<N>_Notes_STUDENT.docx to fill in, and '
            'Day<N>_Notes_KEY.docx with the answers.')
    _b(doc, 'Quiz/. Quiz_STUDENT.docx and Quiz_KEY.docx.')
    _b(doc, 'FRQ/. One free-response practice item per topic, in the College '
            'Board question style, with a separate answer key.')
    _b(doc, 'Teacher_Guide.docx. Objectives with their CED codes, the day-by-day '
            'plan with timings, the traps, differentiation, and the exit ticket '
            'with answers.')

    _h(doc, 'What is where')
    rows = []
    for u in ('1', '2', '3', '4'):
        title, periods, weight, _pc = CED[u]
        days = day_counts.get(u)
        rows.append([f'Unit {u}', title, days if days else 'sold separately',
                     periods, weight])
    _table(doc, ['Unit', 'Title', 'Teaching days', 'CED class periods',
                 'Exam weighting'], rows)
    _p(doc, 'CED class periods are 45 minutes, in College Board\'s own '
            'wording. This kit is built on 60-minute days, so the two columns '
            'are not directly comparable. Course_Resources has the arithmetic.',
       size=9, color=MUTED, space_after=4)

    _h(doc, 'Three conventions to know')
    _b(doc, 'Two editions, one deck. The teacher edition adds speaker notes and '
            'a facilitation callout. Nothing else differs, so either one can be '
            'projected, and the answers to the bell ringer are in the notes '
            'rather than on the slide.')
    _b(doc, 'Every day opens with a bell ringer that retrieves earlier material, '
            'not today\'s. The answers are in the speaker notes and in the '
            'teacher guide.')
    _b(doc, 'The website is where scores are recorded. The printed materials are '
            'not auto-graded; the lesson pages are.')

    _foot(doc, 'Start Here')
    doc.save(path)


def build_how_to_use(path, day_counts):
    doc = Document()
    _title(doc, 'AP COMPUTER SCIENCE A', 'How to Use This Course',
           'The teaching workflow: the three surfaces, the daily rhythm, and '
           'what to do before, during and after a lesson.')

    _h(doc, 'The three surfaces, and what each one is for')
    _table(doc, ['Surface', 'What it is', 'Who touches it'], [
        ['Slides', 'The daily instruction you project. One deck per teaching day.',
         'You, on the projector'],
        ['This folder', 'What you print: guided notes, quizzes, free-response '
         'practice, and every key.', 'You and your students, on paper'],
        ['The website', 'Where students practice and the only place scores are '
         'recorded.', 'Students, signed in'],
    ])

    _h(doc, 'The daily rhythm')
    _b(doc, 'Open with the bell ringer on slide 2. It retrieves earlier '
            'material rather than previewing today, which is the point: '
            'spacing beats review.')
    _b(doc, 'Work the slides. Each teaching segment on the deck matches a '
            'section of the student guided notes, so students fill in as you go.')
    _b(doc, 'Stop at the worked example and take a prediction before running '
            'anything. The wrong prediction is worth more than the right answer.')
    _b(doc, 'Close with the exit ticket in the teacher guide. Answers are '
            'printed there and nowhere the students can see.')

    _h(doc, 'Before, during, after')
    _table(doc, ['When', 'What'], [
        ['Before', 'Read the Teacher_Guide for the topic. The traps section '
         'tells you what the class will get wrong before they do.'],
        ['Before', 'Print the guided notes for the day and the quiz for the '
         'topic. Keys stay on your desk.'],
        ['During', 'Project either edition. Use the speaker notes for bell '
         'ringer answers and facilitation cues.'],
        ['During', 'Differentiation in the teacher guide has three support '
         'moves and four stretch tasks, written for that topic rather than in '
         'general.'],
        ['After', 'Send students to the lesson page for the auto-graded '
         'practice. That is the only surface that reports to your gradebook.'],
        ['After', 'Assign the FRQ packet when the topic is finished. The key '
         'carries the reference solution, hints in the order to give them, and '
         'the sample inputs.'],
    ])

    _h(doc, 'What is auto-graded, and what is not')
    _p(doc, 'The lesson pages carry the checks for understanding, the code '
            'exercises, the debugging exercise and the topic quiz, and all of '
            'them report to your gradebook when a student submits. Everything '
            'in this folder is print material and is graded by you.',
       size=10.5, space_after=4)

    _h(doc, 'A note on the free-response packets')
    _p(doc, 'Each topic ships one item in the College Board question style, '
            'labeled with which of the four exam question types it practices. '
            'The student packet carries the question, the requirements and a '
            'starter. The key is a separate file and says so on its first line: '
            'the reference solution, the hints and the sample inputs are in the '
            'key only. Hidden test cases used by the site\'s auto-grader are '
            'deliberately not printed.',
       size=10.5, space_after=4)

    _foot(doc, 'How to Use This Course')
    doc.save(path)


def build_pacing_full_year(path, day_counts):
    doc = Document()
    _title(doc, 'COURSE RESOURCES', 'Pacing Guide: Full Year',
           'Teaching days per unit, counted from the folders, beside College '
           'Board\'s own class-period bands.')

    _h(doc, 'The two numbers, and why they differ')
    _p(doc, 'This kit is built on 60-minute days. The CED\'s bands are '
            '45-minute periods, in its own wording: "pacing is based on '
            '45-minute class periods, meeting five days each week for a full '
            'academic year." Comparing the columns directly makes this course '
            'look short by about three weeks. Converted to minutes, Units 2 to '
            '4 land inside the CED band and within one percent of its midpoint.',
       size=10.5, space_after=4)

    rows = []
    total_days = 0
    for u in ('1', '2', '3', '4'):
        title, periods, weight, _pc = CED[u]
        days = day_counts.get(u)
        if days:
            total_days += days
        rows.append([f'Unit {u}', title, days if days else 'not in this kit',
                     f'{days * 60} min' if days else '',
                     periods, weight])
    _table(doc, ['Unit', 'Title', 'Teaching days', 'Minutes',
                 'CED periods (45 min)', 'Exam weight'], rows)
    _p(doc, f'Teaching days in this kit: {total_days}. Assessment, review and '
            'mock-exam days are additional and are listed below.',
       size=9.5, color=MUTED, space_after=4)

    _h(doc, 'Progress Checks, as the CED splits them')
    _p(doc, 'These are College Board\'s own groupings, not ours. Each one '
            'is a natural checkpoint.', size=10, color=MUTED, space_after=4)
    pc_rows = []
    for u in ('1', '2', '3', '4'):
        for part, topics in CED[u][3]:
            pc_rows.append([f'Unit {u}', part, topics])
    _table(doc, ['Unit', 'Progress Check', 'Topics'], pc_rows)

    _h(doc, 'Assessment days already decided')
    _b(doc, 'Unit 4 gets two test days: multiple choice first, then free '
            'response on the following day.')
    _b(doc, 'The mock exam runs across two days rather than one long session.')
    _b(doc, 'Checkpoint quizzes are short enough to open a period rather than '
            'consume one.')
    _p(doc, 'A full day-numbered calendar is not printed here. The Year Map '
            'that carries it lives outside this kit, and reproducing it from '
            'memory would put a guess in a table.',
       size=9.5, color=MUTED, space_after=4)

    _foot(doc, 'Pacing Guide, Full Year')
    doc.save(path)


def build_pacing_block(path, day_counts):
    doc = Document()
    _title(doc, 'COURSE RESOURCES', 'Pacing Guide: Block and Semester',
           'The same content on a block schedule, and the one thing that makes '
           'it fit.')

    _h(doc, 'How a block period maps to this kit')
    _p(doc, 'A teaching day in this kit is 60 minutes. A block period is '
            'usually 85 to 95 minutes, so plan roughly three kit days into two '
            'blocks rather than one day per block. The day boundaries in the '
            'decks are natural stopping points, and a block that ends mid-day '
            'is fine as long as the guided notes go home complete.',
       size=10.5, space_after=4)

    rows = []
    for u in ('1', '2', '3', '4'):
        title, _periods, _w, _pc = CED[u]
        days = day_counts.get(u)
        blocks = round(days * 60 / 90) if days else None
        rows.append([f'Unit {u}', title, days if days else 'not in this kit',
                     blocks if blocks else '', f'{days * 60} min' if days else ''])
    _table(doc, ['Unit', 'Title', 'Teaching days (60 min)',
                 'Blocks (90 min, approx)', 'Minutes'], rows)
    _p(doc, 'Block counts are the minute totals divided by 90 and rounded. '
            'They are arithmetic, not a plan: a block that splits a worked '
            'example in half costs more than the rounding saves.',
       size=9.5, color=MUTED, space_after=4)

    _h(doc, 'The one thing that makes the block year fit')
    _p(doc, 'Checkpoint quizzes have to open a block rather than consume one. '
            'Give them their own periods and the block year runs out of float '
            'entirely. That is worth deciding in August rather than discovering '
            'in March.', size=10.5, space_after=4)

    _h(doc, 'Semester course')
    _p(doc, 'A single-semester course cannot cover all four units at this '
            'depth. Units 1 and 2 are the coherent half: they carry 40 to 60 '
            'percent of the exam between them and every later unit depends on '
            'them. Unit 4 is the heaviest unit in the CED and the worst one to '
            'compress.', size=10.5, space_after=4)

    _foot(doc, 'Pacing Guide, Block and Semester')
    doc.save(path)


def build_course_docs(root, day_counts):
    """All four course-level documents. Safe to call once per unit build."""
    build_start_here(os.path.join(root, 'Start_Here.docx'), day_counts)
    build_how_to_use(os.path.join(root, 'How_To_Use_This_Course.docx'), day_counts)
    res = os.path.join(root, 'Course_Resources')
    os.makedirs(res, exist_ok=True)
    build_pacing_full_year(os.path.join(res, 'Pacing_Guide_Full_Year.docx'),
                           day_counts)
    build_pacing_block(os.path.join(res, 'Pacing_Guide_Block_and_Semester.docx'),
                       day_counts)
    return 4
