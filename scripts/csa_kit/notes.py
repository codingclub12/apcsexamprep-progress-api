"""
Guided-notes builder, student and key editions.

Format copied from the Unit 1 packets that already ship
(Day1_Notes_STUDENT.docx): a vocabulary table with the definitions blank, the
day's key ideas truncated mid-sentence so students complete them while the
slide is up, and a closing fill-in-the-blank section where a single term is
removed from each sentence.

The two editions come from one content source. The STUDENT edition blanks the
answers; the KEY edition prints them underlined so a teacher can grade at a
glance. That is why they can never drift apart.

No em-dashes anywhere in generated text.
"""

import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from csa_kit.differentiation import DIFFERENTIATION

BLANK = '_' * 22
NAVY = RGBColor(0x10, 0x24, 0x3A)
ACCENT = RGBColor(0x2C, 0x6B, 0xAF)
MUTED = RGBColor(0x5A, 0x6B, 0x7B)

TRADEMARK = ("AP is a trademark of the College Board, which was not involved in "
             "the production of, and does not endorse, this resource.")

# How much of a key-idea sentence the student sees before the blank starts.
# This is a CEILING measured in characters, not the cut point. See _cloze_head.
TRUNCATE_AT = 58

# Never hand back a stub so short the sentence has stopped meaning anything.
# If honoring the word boundary would cut below this, the hard character cut
# is used instead, because a slightly ugly line beats a meaningless one.
TRUNCATE_FLOOR = 28


def _p(doc, text='', size=11, bold=False, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.name = 'Calibri'
    if color is not None:
        r.font.color.rgb = color
    return p


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.name = 'Cambria'
    r.font.color.rgb = NAVY
    return p


def _cloze_head(sentence):
    """The visible part of a cloze line, cut at a word boundary.

    Until 2026-09-04 this was a bare sentence[:TRUNCATE_AT], which cuts on a
    character count and therefore lands wherever it lands. Real lines it
    produced, from the shipped Unit 2 student notes:

        "...evaluates to exactly one of two value"        (lost the s)
        "...combine booleans: && means and, || m"         (mid-word)
        "A truth table lists every possible combination of the inpu"

    A student reading "the inpu" cannot tell whether the word was cut or
    misspelled, and the sheet reads as broken rather than as a fill-in.

    So: cut at the last space at or before the ceiling, then drop trailing
    punctuation that would dangle in front of the blank. Fall back to the hard
    cut when a word boundary would leave too little to be worth reading, which
    is what TRUNCATE_FLOOR is for.
    """
    if len(sentence) <= TRUNCATE_AT:
        return sentence.rstrip()
    window = sentence[:TRUNCATE_AT + 1]
    space = window.rfind(' ')
    head = window[:space] if space >= TRUNCATE_FLOOR else sentence[:TRUNCATE_AT]
    return head.rstrip().rstrip(',;:.')


def _cloze_line(doc, sentence, key_edition):
    """A key idea, truncated so the student finishes writing it."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    if key_edition:
        r = p.add_run(sentence)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
    else:
        head = _cloze_head(sentence)
        r = p.add_run(head)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        r2 = p.add_run('  ' + BLANK)
        r2.font.size = Pt(11)
        r2.font.name = 'Calibri'
    return p


def _fill_blank(doc, template, answer, key_edition):
    """template contains {} where the term is removed."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    parts = template.split('{}')
    for i, part in enumerate(parts):
        r = p.add_run(part)
        r.font.size = Pt(11)
        r.font.name = 'Calibri'
        if i < len(parts) - 1:
            a = p.add_run(answer if key_edition else '_' * 14)
            a.font.size = Pt(11)
            a.font.name = 'Calibri'
            a.font.underline = True
            if key_edition:
                a.font.bold = True
                a.font.color.rgb = ACCENT
    return p


def build_notes(path, topic, title, day, handle, sections, vocab, blanks, key_edition):
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)

    _p(doc, 'GUIDED NOTES' + ('  |  ANSWER KEY' if key_edition else ''),
       size=10, bold=True, color=ACCENT, space_after=2)
    _p(doc, f'Topic {topic}: {title}', size=16, bold=True, color=NAVY, space_after=2)
    _p(doc, f'Day {day}  |  Name: ' + '_' * 26, size=11, space_after=4)
    _p(doc, f'Practice online: apcsexamprep.com/pages/{handle}   Your work there is scored automatically.',
       size=9.5, color=MUTED, space_after=10)

    # ── vocabulary ───────────────────────────────────────────────────────────
    _heading(doc, 'Vocabulary')
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ('Term', 'Definition' if key_edition else 'Definition (fill in)')):
        cell.text = ''
        r = cell.paragraphs[0].add_run(text)
        r.font.bold = True
        r.font.size = Pt(10.5)
        r.font.name = 'Calibri'
    for term, definition in vocab:
        cells = table.add_row().cells
        r = cells[0].paragraphs[0].add_run(term)
        r.font.size = Pt(10.5)
        r.font.name = 'Calibri'
        r.font.bold = True
        if key_edition:
            r2 = cells[1].paragraphs[0].add_run(definition)
            r2.font.size = Pt(10)
            r2.font.name = 'Calibri'
        else:
            cells[1].text = ''

    # ── the day's key ideas, truncated ───────────────────────────────────────
    for name, ideas in sections:
        _heading(doc, name)
        for idea in ideas:
            _cloze_line(doc, idea, key_edition)

    # ── fill in the blanks ───────────────────────────────────────────────────
    _heading(doc, 'Fill in the blanks')
    for template, answer in blanks:
        _fill_blank(doc, template, answer, key_edition)

    _p(doc, '', space_after=10)
    _p(doc, TRADEMARK, size=7.5, color=MUTED, space_after=2)
    _p(doc, f'APCSExamPrep.com   Topic {topic} day {day}', size=8, color=MUTED)
    doc.save(path)


def build_quiz(path, topic, title, handle, questions, key_edition):
    """Short end-of-topic quiz, student and key editions. questions is a list of
    dicts: {stem, options[], answer_index, why}."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)

    _p(doc, 'TOPIC QUIZ' + ('  |  ANSWER KEY' if key_edition else ''),
       size=10, bold=True, color=ACCENT, space_after=2)
    _p(doc, f'Topic {topic}: {title}', size=16, bold=True, color=NAVY, space_after=2)
    if not key_edition:
        _p(doc, 'Name: ' + '_' * 30 + '     Score: ____ / ' + str(len(questions)),
           size=11, space_after=4)
    _p(doc, f'The same questions are auto-graded at apcsexamprep.com/pages/{handle}',
       size=9.5, color=MUTED, space_after=12)

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'{i}. {q["stem"]}')
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.name = 'Calibri'
        for j, opt in enumerate(q['options']):
            letter = 'ABCD'[j]
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Inches(0.35)
            op.paragraph_format.space_after = Pt(2)
            run = op.add_run(f'{letter}. {opt}')
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'
            if key_edition and j == q['answer_index']:
                run.font.bold = True
                run.font.color.rgb = ACCENT
        if key_edition:
            wp = doc.add_paragraph()
            wp.paragraph_format.left_indent = Inches(0.35)
            wp.paragraph_format.space_after = Pt(12)
            wr = wp.add_run('Why: ' + q['why'])
            wr.font.size = Pt(9.5)
            wr.font.italic = True
            wr.font.name = 'Calibri'
            wr.font.color.rgb = MUTED
        else:
            _p(doc, '', space_after=8)

    _p(doc, TRADEMARK, size=7.5, color=MUTED, space_after=2)
    _p(doc, f'APCSExamPrep.com   Topic {topic} quiz', size=8, color=MUTED)
    doc.save(path)


STRUCTURAL = ('bell ringer', 'objectives and guided', 'worked', 'misconception',
              'stop and think', 'guided practice', 'independent practice',
              'exit ticket')


def _is_structural(label):
    """True for a schedule segment whose content comes from a named field
    (warmup, worked, misconception, discussion) rather than from sections."""
    low = label.lower()
    return any(low.startswith(p) for p in STRUCTURAL)


def _bullet(doc, text, size=10.5, indent=0.25):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    return p


def _mono(doc, text, indent=0.45):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.size = Pt(10)
        r.font.name = 'Consolas'
        r.font.color.rgb = NAVY
    return doc


def _segment(doc, label, minutes):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.name = 'Calibri'
    r.font.color.rgb = NAVY
    if minutes:
        r2 = p.add_run(f'   ({minutes} min)')
        r2.font.size = Pt(9.5)
        r2.font.name = 'Calibri'
        r2.font.color.rgb = MUTED
    return p



def build_frq(path, item, key_edition):
    """One free-response practice item, student packet or answer key.

    The 53 items are authored in seed/csa-frq and already drive the live
    auto-graded pages. Nothing here is new content: this renders what exists
    into the printable packet the kit was missing, the same way the teacher
    guide renders sections the deck was already using.

    THE KEY IS A SEPARATE FILE AND SAYS SO ON ITS FIRST LINE. The reference
    solution, the hints and the sample inputs all live in the key. A student
    packet that quietly carried the solution would be the bell-ringer bug
    again, one surface further out.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.7)
        sec.left_margin = sec.right_margin = Inches(0.8)

    label = 'ANSWER KEY' if key_edition else 'FREE RESPONSE PRACTICE'
    _p(doc, label, size=10, bold=True, color=ACCENT, space_after=2)
    _p(doc, f"Topic {item['lesson']}: {item['name']}", size=19, bold=True,
       color=NAVY, space_after=2)
    _p(doc, f"{item['title']}   |   Question type: {item['frqTypeLabel']}",
       size=9.5, color=MUTED, space_after=10)

    if key_edition:
        _p(doc, 'Teacher copy. Do not hand this page to students.',
           size=9.5, color=MUTED, space_after=8)

    if item.get('brief'):
        _heading(doc, 'Why this question')
        _p(doc, item['brief'], size=10.5, space_after=6)

    if item.get('given'):
        _heading(doc, 'Given')
        _p(doc, item['given'], size=10.5, space_after=6)

    if item.get('parts'):
        _heading(doc, 'Write the following')
        for part in item['parts']:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(part['label'] + '  ')
            r.font.size = Pt(10.5)
            r.font.bold = True
            r.font.name = 'Calibri'
            r2 = p.add_run(part['text'])
            r2.font.size = Pt(10.5)
            r2.font.name = 'Calibri'

    if item.get('task'):
        _heading(doc, 'Requirements')
        for line in item['task']:
            _bullet(doc, line)

    if item.get('starter') and not key_edition:
        _heading(doc, 'Starter')
        _mono(doc, item['starter'], indent=0.2)

    if key_edition:
        if item.get('reference'):
            _heading(doc, 'Reference solution')
            _mono(doc, item['reference'], indent=0.2)
        if item.get('hints'):
            _heading(doc, 'Hints, in the order to give them')
            for h in item['hints']:
                _bullet(doc, h)
        if item.get('sampleCases'):
            _heading(doc, 'Sample inputs')
            _p(doc, 'Visible cases only. The hidden cases used by the '
                    'auto-grader are deliberately not printed.',
               size=9.5, color=MUTED, space_after=4)
            for c in item['sampleCases']:
                _mono(doc, (c.get('stdin') or '').rstrip('\n'), indent=0.2)

    _p(doc)
    _p(doc, TRADEMARK, size=7.5, color=MUTED, space_after=0)
    _p(doc, f"APCSExamPrep.com   Topic {item['lesson']} free response",
       size=7.5, color=MUTED, space_after=0)
    doc.save(path)

def build_teacher_guide(path, topic, title, handle, days, vocab, quiz, graded_line):
    """The full teacher guide for one topic.

    THIS REPLACED A ONE PAGE LESSON MAP ON 2026-09-04, and the reason is worth
    keeping. The Unit 1 guides Tanner already ships carry the learning
    objectives with their CED codes, every teaching segment's actual content,
    the bell ringer WITH its debrief answers, an exit ticket with answers, and
    the traps the topic sets. The generated Units 2-4 guide carried two timing
    tables and four sentences, so a teacher opening 1.3 and then 2.3 met two
    different products.

    Nothing here is newly authored. Every section below is rendered from the
    same topic dict that already drives the slides, the guided notes and the
    quiz, which is the kit's founding rule: one content source per topic, so
    the surfaces cannot drift. The content was there the whole time and only
    the deck and the notes were reading it.

    ONE SECTION OF THE UNIT 1 GUIDE IS DELIBERATELY ABSENT. Unit 1 carries
    Differentiation, three Support moves and four Stretch tasks per topic, and
    there is nothing in the content dict to render that from. Inventing it here
    would be authoring 266 items nobody has reviewed, which this repo treats as
    a human decision rather than a generator feature.
    """
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)

    _p(doc, 'TEACHER GUIDE', size=10, bold=True, color=ACCENT, space_after=2)
    _p(doc, f'Topic {topic}: {title}', size=19, bold=True, color=NAVY, space_after=2)
    _p(doc, f'{len(days)} days   |   Prepared for the May 2027 AP CSA exam '
            '(2025 Course and Exam Description, four-unit structure).',
       size=9.5, color=MUTED, space_after=10)

    # ── learning objectives ──────────────────────────────────────────────────
    seen = []
    for d in days:
        for text, code in d.get('objectives', []):
            if code not in [c for _, c in seen]:
                seen.append((text, code))
    if seen:
        _heading(doc, 'Learning objectives')
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = 'Code', 'Objective'
        for text, code in seen:
            row = table.add_row().cells
            row[0].text = code.replace('LO ', '')
            row[1].text = text
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
                        r.font.name = 'Calibri'

        _heading(doc, 'Students will be able to')
        for text, _code in seen:
            _bullet(doc, text)

    # ── how the days run ─────────────────────────────────────────────────────
    _heading(doc, 'How the days run')
    for d in days:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"Day {d['day']}")
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.name = 'Cambria'
        r.font.color.rgb = ACCENT
        if d.get('focus'):
            _p(doc, d['focus'], size=10, color=MUTED, space_after=4)

        # Attach the day's sections to its teaching segments BY ORDER.
        #
        # Two earlier attempts failed and both are worth naming. Matching the
        # schedule label to the section name exactly missed almost everything,
        # because a label is a description and a name is a title: "Why order
        # matters in a chain" against "Chains and order". Normalizing
        # punctuation fixed one case out of 114 and left 76 teaching segments
        # printing a heading and a timing with no content beneath, which is the
        # most useful thing the guide carries.
        #
        # The order does line up, and it lines up in both shapes the content
        # takes. Measured across all 76 days: 38 have one section per teaching
        # segment, and 38 have one long segment covering two sections. So walk
        # the segments in order, hand each the next section, and let the LAST
        # teaching segment absorb whatever is left. Both shapes come out right
        # and nothing is silently dropped.
        sections = list(d.get('sections', []))
        teach_idx = [i for i, (_m, lab) in enumerate(d.get('schedule', []))
                     if not _is_structural(lab)]
        attached = {}
        if teach_idx:
            per = len(sections) // len(teach_idx)
            per = max(per, 1)
            cursor = 0
            for n, i in enumerate(teach_idx):
                last = (n == len(teach_idx) - 1)
                take = sections[cursor:] if last else sections[cursor:cursor + per]
                attached[i] = take
                cursor += len(take)
        warm = d.get('warmup') or []
        for seg_i, (minutes, label) in enumerate(d.get('schedule', [])):
            _segment(doc, label, minutes)
            low = label.lower()
            if low.startswith('bell ringer') and len(warm) >= 3:
                _bullet(doc, warm[1])
                if len(warm) > 3 and warm[3]:
                    _mono(doc, warm[3])
                _bullet(doc, f'Debrief: {warm[2]}')
            elif seg_i in attached:
                for name, ideas in attached[seg_i]:
                    if len(attached[seg_i]) > 1:
                        _p(doc, name, size=10.5, bold=True, space_after=2)
                    for idea in ideas:
                        _bullet(doc, idea)
            elif low.startswith('worked'):
                w = d.get('worked') or {}
                if w.get('note'):
                    _bullet(doc, w['note'])
                elif w.get('caption'):
                    _bullet(doc, w['caption'])
            elif low.startswith('misconception'):
                m = d.get('misconception') or {}
                if m.get('truth'):
                    _bullet(doc, f"{m.get('heading', 'Misconception')}: {m['truth']}")
            elif low.startswith('stop and think'):
                for q in (d.get('discussion') or []):
                    _bullet(doc, q)
                if d.get('extra'):
                    _bullet(doc, f"Homework: {d['extra']}")

    # ── exit ticket ──────────────────────────────────────────────────────────
    if quiz:
        _heading(doc, 'Exit ticket, with answers')
        _p(doc, 'Answers are shown. Do not hand this page to students.',
           size=9.5, color=MUTED, space_after=6)
        for i, q in enumerate(quiz, 1):
            _p(doc, f"{i}. {q['stem']}", size=10.5, space_after=2)
            for j, opt in enumerate(q['options']):
                mark = '  <-- answer' if j == q['answer_index'] else ''
                _bullet(doc, f"{chr(65 + j)}. {opt}{mark}", size=10, indent=0.45)
            if q.get('why'):
                _p(doc, f"Why: {q['why']}", size=9.5, color=MUTED, space_after=8)

    # ── traps ────────────────────────────────────────────────────────────────
    traps = []
    for d in days:
        m = d.get('misconception') or {}
        if m.get('truth'):
            traps.append((m.get('heading', 'Misconception'), m['truth']))
        b = d.get('break_it') or {}
        if b.get('why'):
            change = (b.get('change') or 'One change').rstrip('.')
            traps.append((f'Change it and see: {change}', b['why']))
    if traps:
        _heading(doc, 'Traps this topic sets')
        for head, body in traps:
            # rstrip the period: several headings already end with one, and
            # "same conditions.. Separate ifs" is how that looked before.
            _p(doc, f'{head.rstrip(".")}. {body}', size=10.5, space_after=6)

    # ── vocabulary ───────────────────────────────────────────────────────────
    if vocab:
        _heading(doc, 'Vocabulary')
        for term, definition in vocab:
            _bullet(doc, f'{term}: {definition}')

    # ── differentiation ──────────────────────────────────────────────────────
    # The one section of the Unit 1 guides that is NOT rendered from the topic
    # dict, because nothing in the dict describes it. It is authored separately
    # in csa_kit/differentiation.py and is the only newly written content in
    # this kit, which is why a topic missing from that file simply omits the
    # section rather than printing a placeholder: a guide with no
    # Differentiation heading is honest, and a heading over invented filler is
    # not.
    diff = DIFFERENTIATION.get(topic)
    if diff:
        _heading(doc, 'Differentiation')
        _p(doc, 'Support', size=11.5, bold=True, color=ACCENT, space_after=3)
        for item in diff['support']:
            _bullet(doc, item)
        _p(doc, 'Stretch', size=11.5, bold=True, color=ACCENT, space_after=3)
        for item in diff['stretch']:
            _bullet(doc, item)

    # ── the three surfaces ───────────────────────────────────────────────────
    _heading(doc, 'On the website (students)')
    _p(doc, 'Everything on the lesson page is auto-graded and reports to your '
            'gradebook the moment a student submits.', size=10.5, space_after=4)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Grid Accent 1'
    for k, v in (('Lesson page', f'apcsexamprep.com/pages/{handle}'),
                 ('Auto-graded on that page', graded_line),
                 ('Your gradebook', 'apcsexamprep.com/pages/cyber-dashboard')):
        row = table.add_row().cells
        row[0].text, row[1].text = k, v
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Calibri'

    _p(doc)
    _p(doc, TRADEMARK, size=7.5, color=MUTED, space_after=0)
    _p(doc, f'APCSExamPrep.com   Topic {topic} teacher guide', size=7.5,
       color=MUTED, space_after=0)
    doc.save(path)
